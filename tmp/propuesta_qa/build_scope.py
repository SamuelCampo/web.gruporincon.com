from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT='/Users/user/projects/projects/Web/web.gruporincon.com/outputs/019ff790-5100-7c10-aaa2-4602911ca668/Alcance_Detallado_MVP_App_Cuotas.docx'
NAVY='0B2A4A'; BLUE='1769AA'; PALE='EAF2F8'; GREEN='DFF2E1'; AMBER='FFF0CC'; RED='FEE4E2'; LINE='D0D5DD'
d=Document(); sec=d.sections[0]; sec.top_margin=Inches(.65); sec.bottom_margin=Inches(.65); sec.left_margin=Inches(.75); sec.right_margin=Inches(.75)
styles=d.styles; styles['Normal'].font.name='Arial'; styles['Normal'].font.size=Pt(10); styles['Normal'].paragraph_format.space_after=Pt(6)
for n,size,color in [('Title',28,NAVY),('Heading 1',18,NAVY),('Heading 2',13,BLUE),('Heading 3',11,NAVY)]:
 s=styles[n]; s.font.name='Arial'; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=True

def shade(cell,color):
 tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),color); tcPr.append(shd)
def margins(cell,top=100,start=120,bottom=100,end=120):
 tc=cell._tc.get_or_add_tcPr(); m=tc.first_child_found_in('w:tcMar')
 if m is None: m=OxmlElement('w:tcMar'); tc.append(m)
 for tag,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
  x=OxmlElement('w:'+tag); x.set(qn('w:w'),str(val)); x.set(qn('w:type'),'dxa'); m.append(x)
def set_cell(cell,text,bold=False,color='000000',fill=None,align=None):
 cell.text=''; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(text); r.bold=bold; r.font.name='Arial'; r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(color)
 if align: p.alignment=align
 cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)
 if fill: shade(cell,fill)
def table(headers,rows,widths=None):
 t=d.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
 for i,h in enumerate(headers): set_cell(t.rows[0].cells[i],h,True,'FFFFFF',BLUE,WD_ALIGN_PARAGRAPH.CENTER)
 for row in rows:
  cells=t.add_row().cells
  for i,v in enumerate(row): set_cell(cells[i],str(v),False,'000000','F8FAFC' if len(t.rows)%2==0 else 'FFFFFF')
 if widths:
  for row in t.rows:
   for i,w in enumerate(widths): row.cells[i].width=Inches(w)
 t.rows[0]._tr.get_or_add_trPr().append(OxmlElement('w:tblHeader'))
 return t
def bullet(text,level=0):
 p=d.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2'); p.add_run(text); return p
def callout(title,text,color=PALE):
 t=d.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; c=t.cell(0,0); shade(c,color); margins(c,160,180,160,180); p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(3); r=p.add_run(title+'\n'); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY); r.font.size=Pt(11); p.add_run(text); d.add_paragraph().paragraph_format.space_after=Pt(0)
def page(): d.add_page_break()

p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(70); r=p.add_run('PROPUESTA DE ALCANCE'); r.bold=True; r.font.size=Pt(30); r.font.color.rgb=RGBColor.from_string(NAVY)
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('MVP de aplicación de comercio a cuotas'); r.font.size=Pt(18); r.font.color.rgb=RGBColor.from_string(BLUE)
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('App móvil multiplataforma + Supabase + panel administrativo básico').italic=True
d.add_paragraph(); callout('Objetivo del documento','Definir exactamente qué se desarrollará por USD 900, qué lógica debe aprobar el cliente y cuáles funciones quedan simplificadas o reservadas para evolución.',PALE)
table(['Dato','Definición'],[['Versión','MVP 1.0'],['Fecha','12 de agosto de 2026'],['Inversión inicial','USD 900'],['Plazo estimado','3 semanas desde anticipo y aprobación'],['Mantenimiento','USD 100/mes; permanencia mínima sugerida de 6 meses'],['Estado','Para revisión y aprobación del cliente']],[2,4.6])
d.add_paragraph(); p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Documento confidencial - Grupo Rincón').font.color.rgb=RGBColor.from_string('667085')
page()
d.add_heading('1. Resumen ejecutivo',level=1)
d.add_paragraph('La primera versión permitirá publicar productos, vender al contado o mediante un plan fijo de cuotas, recibir comprobantes de pago y controlar manualmente la aprobación, el despacho y el avance de las cuotas. La solución prioriza una salida rápida y un costo inicial bajo, sin construir todavía un marketplace multi-vendedor completo ni automatizaciones financieras de alto costo.')
callout('Decisión técnica propuesta','React Native con Expo para la app móvil y Supabase para autenticación, base de datos PostgreSQL, almacenamiento y reglas de seguridad. Esta combinación reduce tiempo de construcción y conserva una ruta clara de escalamiento.',GREEN)
d.add_heading('2. Objetivos del MVP',level=1)
for x in ['Validar comercialmente la venta de productos a crédito con una inicial fija del 45%.','Permitir que el comprador elija pago semanal o quincenal bajo reglas simples y auditables.','Centralizar productos, pedidos, pagos, tasa de cambio y estados operativos.','Reducir la dependencia de hojas de cálculo sin intentar reemplazar en esta fase un ERP o sistema contable.','Dejar una base técnica escalable para automatizaciones y operación multi-tienda.']: bullet(x)
d.add_heading('3. Usuarios y roles',level=1)
table(['Rol','Capacidades en MVP'],[['Comprador','Registro, inicio de sesión, catálogo, detalle, selección de modalidad, carga de comprobante, consulta de pedido y cronograma.'],['Administrador','Gestión básica de productos, tasa USDT, pedidos, revisión manual de comprobantes, aprobación/rechazo, estados de entrega y registro de cuotas.'],['Vendedor/Tienda','Se representa visualmente en catálogo cuando aplique, pero no tiene panel autogestionado independiente en esta fase.']],[1.5,5.1])
page()
d.add_heading('4. Alcance funcional detallado',level=1)
sections=[
('4.1 Acceso y perfil',['Registro con nombre, correo, teléfono y contraseña.','Inicio y cierre de sesión; recuperación básica de contraseña.','Edición de datos esenciales del perfil.','Aceptación de términos y política de privacidad mediante enlaces suministrados por el cliente.']),
('4.2 Catálogo y productos',['Pantalla inicial con categorías, tiendas destacadas y productos.','Búsqueda básica y navegación por categoría.','Detalle con imágenes, descripción, precio a cuotas, precio de contado y opciones disponibles.','Carga y edición de productos desde el panel administrativo.','Campos de precio, rebaja de contado y zona de delivery gratis.']),
('4.3 Compra y financiamiento',['Elección entre contado y cuotas.','Inicial obligatoria calculada en 45%.','Saldo financiado calculado en 55%.','Elección de 4 pagos semanales o 2 pagos quincenales.','Generación automática de montos y fechas del cronograma.','Visualización de pagado, pendiente y próximo vencimiento.']),
('4.4 Pagos y moneda',['Carga de comprobante de Pago Móvil o transferencia.','Aprobación o rechazo manual por el administrador.','Tasa USDT P2P registrada manualmente por el administrador.','Cálculo referencial en moneda local y conservación de la tasa aplicada a cada transacción.','Sin cobro automático, débito recurrente ni conciliación bancaria.']),
('4.5 Pedidos y despacho',['Creación de pedido con resumen financiero.','Estados mínimos: pendiente de pago, pago en revisión, aprobado, en proceso, enviado, entregado y cancelado.','El despacho solo se habilita después de aprobar la inicial del 45%.','Cambios de estado realizados manualmente por el administrador.']),
('4.6 Panel administrativo',['Inicio de sesión administrativo.','Resumen básico de pedidos y pagos pendientes.','CRUD básico de productos y categorías.','Configuración manual de tasa USDT.','Revisión de comprobantes y registro de cuotas pagadas.','Consulta de cliente, producto, saldo y estado de entrega.'])]
for title,items in sections:
 d.add_heading(title,level=2)
 for x in items: bullet(x)
page()
d.add_heading('5. Lógica de negocio para aprobación',level=1)
table(['Regla','Definición acordada','Resultado'],[
['Inicial','Precio a cuotas × 45%','Pago mínimo para autorizar despacho'],['Saldo','Precio a cuotas - inicial','55% financiado'],['Semanal','Saldo ÷ 4','4 cuotas iguales con vencimientos semanales'],['Quincenal','Saldo ÷ 2','2 cuotas iguales con vencimientos quincenales'],['Contado','Precio a cuotas - rebaja configurable','Precio final de pago inmediato'],['Moneda local','Monto USD × tasa USDT registrada','Referencia local conservando la tasa usada'],['Aprobación','Revisión manual del comprobante','Aprobado o rechazado'],['Despacho','Inicial aprobada','Habilita avance a proceso/envío/entrega']],[1.3,3.6,1.8])
callout('Ejemplo','Para un producto de USD 70: inicial USD 31,50; saldo USD 38,50; cuatro cuotas semanales de USD 9,625 (presentadas según regla de redondeo acordada) o dos quincenales de USD 19,25. Con rebaja de USD 15, el contado es USD 55.',AMBER)
d.add_heading('Decisiones que el cliente debe confirmar',level=2)
for x in ['¿El redondeo se hará a 2 decimales y el ajuste de centavos se aplicará en la última cuota?','¿La primera cuota vence 7 o 15 días después de aprobar la inicial, según frecuencia?','¿La tasa USDT aplicable es la registrada al generar cada pago o al aprobarlo? Recomendación: conservar la tasa al generar/reporta el pago.','¿Se permite pago anticipado de cuotas sin penalización? Recomendación: sí.','¿Qué sucede con retrasos? El MVP registra vencido, pero no calcula mora ni penalidad.','¿La rebaja de contado es fija por producto y solo aplica a pago inmediato? Recomendación: sí.']:
 bullet(x)
d.add_heading('Estados propuestos',level=2)
table(['Entidad','Estados'],[['Pago','Reportado → En revisión → Aprobado / Rechazado'],['Pedido','Pendiente de inicial → En proceso → Enviado → Entregado / Cancelado'],['Cuota','Pendiente → Reportada → Pagada / Vencida']],[1.5,5.2])
page()
d.add_heading('6. Entregables',level=1)
table(['#','Entregable','Criterio de aceptación'],[
['1','Código fuente de la app móvil','Compila desde el repositorio entregado y cubre los flujos aprobados.'],['2','Proyecto Supabase','Esquema, autenticación, almacenamiento y políticas básicas configuradas.'],['3','Panel administrativo básico','Permite operar productos, tasa, pedidos, pagos y estados incluidos.'],['4','Base de datos inicial','Tablas y relaciones para usuarios, productos, pedidos, pagos, cuotas y tasa.'],['5','Despliegue de prueba','Ambiente funcional para validación del cliente.'],['6','Documentación breve','Accesos, configuración, flujo de operación y respaldo.'],['7','Capacitación','Una sesión remota de hasta 60 minutos para el equipo administrador.'],['8','Garantía correctiva','15 días posteriores a la aceptación para errores del alcance aprobado.']],[.4,2.2,4.1])
d.add_heading('7. Criterios de aceptación',level=1)
for x in ['Un comprador puede registrarse e iniciar sesión.','El catálogo muestra productos creados por administración.','El sistema calcula correctamente 45%, 55%, planes semanal/quincenal y contado.','El comprador puede cargar un comprobante.','El administrador puede aprobarlo o rechazarlo.','Un pedido no puede avanzar a despacho sin inicial aprobada.','El administrador puede registrar cuotas y consultar saldo/estado.','La app y el panel funcionan en el ambiente de prueba acordado.']: bullet(x)
d.add_heading('8. Datos y materiales requeridos al cliente',level=1)
for x in ['Nombre comercial, logotipo y paleta de marca.','Listado inicial de productos, categorías, imágenes, precios, rebajas y zonas de delivery.','Textos legales: términos, privacidad y políticas de crédito/cobro.','Datos de pago y procedimiento interno de validación.','Usuario responsable de aprobar productos, pagos y pruebas.','Cuentas de Apple/Google si se contrata publicación en tiendas.']: bullet(x)
page()
d.add_heading('9. Alcance simplificado y exclusiones',level=1)
table(['Clasificación','Elemento','Tratamiento'],[
['Simplificado','KYC / identidad','Captura o carga de información para revisión manual; sin biometría ni proveedor KYC.'],['Simplificado','Multi-tienda','Visualización de tiendas con operación centralizada; sin portal independiente por vendedor.'],['Simplificado','Notificaciones','Avisos básicos dentro del producto; WhatsApp, SMS y push avanzado quedan para evolución.'],['Excluido','Pasarela de pago','No incluye cobro automático, tokenización, débito ni conciliación.'],['Excluido','Mora e interés','No calcula intereses, penalidades, refinanciamiento ni scoring crediticio.'],['Excluido','Comisiones/liquidación','No incluye comisión de plataforma ni liquidaciones a vendedores.'],['Excluido','Logística integrada','No genera etiquetas ni integra transportistas o tracking externo.'],['Excluido','Reportes avanzados','No incluye analítica, exportación masiva, contabilidad ni ERP.'],['Excluido','Publicación y terceros','Tasas, licencias y consumos de Apple, Google, Supabase, dominios, correo, KYC o mensajería.']],[1.25,2,3.45])
callout('Control de cambio','Toda función no descrita como incluida se registra como solicitud de cambio. Puede cotizarse por separado o priorizarse dentro de la bolsa mensual de evolución.',RED)
d.add_heading('10. Plan de trabajo y tiempos',level=1)
table(['Semana','Objetivo','Hitos'],[['1','Definición y base técnica','Aprobación de lógica; configuración Supabase; autenticación; catálogo base.'],['2','Operación comercial','Detalle, cálculo, checkout, comprobantes, cronograma y panel.'],['3','Cierre','Estados, pruebas, correcciones, despliegue, documentación y capacitación.']],[.8,2,4])
d.add_paragraph('El plazo depende de recibir contenidos y aprobaciones sin demoras. Cambios de lógica posteriores a la aprobación pueden mover la fecha de entrega.')
d.add_heading('11. Inversión y condiciones',level=1)
table(['Concepto','Monto','Condición'],[['Desarrollo e implementación MVP 1.0','USD 900','50% al iniciar / 50% contra entrega y despliegue'],['Mantenimiento y evolución','USD 100/mes','Contrato mínimo sugerido: 6 meses'],['Servicios de terceros','No incluidos','Pagados directamente por el cliente según consumo']],[2.4,1.4,3])
d.add_heading('Mantenimiento mensual',level=2)
for x in ['Monitoreo básico de la infraestructura y revisión de respaldos.','Corrección de fallas reproducibles del producto en operación.','Hasta 4 horas mensuales de evolución, acumulables por un máximo de 3 meses.','Priorización mensual de mejoras por acuerdo entre las partes.','No incluye rediseños completos, nuevas integraciones complejas ni soporte operativo 24/7.']: bullet(x)
d.add_heading('12. Riesgos y responsabilidades',level=1)
table(['Riesgo / dependencia','Tratamiento'],[['Aprobación de tiendas','Apple y Google pueden solicitar ajustes o rechazar publicaciones; no depende exclusivamente del proveedor.'],['Tasa USDT','La fuente y el momento de aplicación deben ser definidos por el cliente; el admin conserva responsabilidad operativa.'],['Cumplimiento de crédito','El cliente define contratos, mora, cobranza, privacidad y obligaciones legales.'],['Calidad de contenidos','El cliente suministra textos, precios e imágenes con derechos de uso.'],['Crecimiento','Volumen, automatizaciones o múltiples vendedores pueden requerir ampliación de arquitectura y presupuesto.']],[2.2,4.6])
page()
d.add_heading('13. Aprobación',level=1)
d.add_paragraph('Con la firma o confirmación escrita, el cliente aprueba la lógica descrita, el alcance incluido, las simplificaciones, exclusiones, inversión y condiciones de entrega.')
table(['Campo','Completar'],[['Nombre y cargo',''],['Empresa',''],['Aprobación','☐ Aprobado   ☐ Aprobado con observaciones   ☐ Requiere cambios'],['Firma',''],['Fecha',''],['Observaciones','']],[2,4.8])

for section in d.sections:
 footer=section.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; run=footer.add_run('Grupo Rincón | Alcance MVP App de Cuotas | Confidencial'); run.font.name='Arial'; run.font.size=Pt(8); run.font.color.rgb=RGBColor.from_string('667085')
d.core_properties.title='Alcance detallado MVP App de Comercio a Cuotas'; d.core_properties.subject='Propuesta técnica y funcional'; d.core_properties.author='Grupo Rincón'
d.save(OUT); print(OUT)
